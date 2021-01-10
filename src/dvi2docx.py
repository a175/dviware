import docx
import sys
import dviware
import os

class DviDocxStackMachine(dviware.DviStackMachine):
    def __init__(self,mathimagebasename,debugmode=False):
        super().__init__(debugmode)
        self.document = docx.Document()
        self.p = None
        self.r = None
        self.is_mathmode = 0
        self.is_display = False
        self.num_of_math = 0
        self.math_image_base_name = mathimagebasename
        self.list_stack=[]
        self.spaceque=None
        self.p_is_empty=False

    def string2length(self,string):
        if string.endswith("pt"):
            return docx.shared.Pt(float(string[:-2]))

    def add_new_run_for_text(self):
        self.r=self.p.add_run()
        if self.stackmemory.f!=-1:
            s=self.fontregister.get_scaledsize(self.stackmemory.f)
            self.r.font.size=docx.shared.Mm(self.get_dimension_as_float(s)/10000000)            
            self.r.font.italic=self.fontregister.is_italic(self.stackmemory.f)
            self.r.font.bold=self.fontregister.is_bold(self.stackmemory.f)
            self.r.font.small_caps=self.fontregister.is_small_caps(self.stackmemory.f)
            self.r.font.name=self.fontregister.get_name(self.stackmemory.f)

    def add_new_paragraph(self,param=None):
        if not self.p_is_empty:
            self.p=self.document.add_paragraph()
            if param:
                if param.startswith("parindent="):
                    indent=self.string2length(param[10:])
                    self.p.paragraph_format.first_line_indent=indent
            self.add_new_run_for_text()
            self.p_is_empty=True

    def add_single_space(self,space=" "):
        if self.spaceque==None:
            self.spaceque=space
    
    def add_mathimage(self):
        self.p_is_empty=False
    
        self.num_of_math = self.num_of_math+1
        if not self.is_display:
            if self.spaceque!=None:
                self.r.add_text(self.spaceque)
                self.spaceque=None
            self.r=self.p.add_run()
            self.r.add_picture(self.math_image_base_name+str(self.num_of_math)+".png")
        else:
            self.r=self.p.add_run()
            self.r.add_break()
            self.r.add_text("      ")
            self.r.add_picture(self.math_image_base_name+str(self.num_of_math)+".png")
            self.r.add_break()
        self.add_new_run_for_text()
            

    def bop(self,cc,p,bb):
        """
        function for DVI.bop.
        """
        ans=super().bop(cc,p,bb)
        if self.p:
            self.document.add_page_break()
        if self.is_mathmode > 0:
            self.add_mathimage()

    def draw_char(self,h,v,c):
        """
        add a character to paragraph as new run.
        """
        ans=super().draw_char(h,v,c)
        self.p_is_empty=False
        if self.spaceque!=None:
            self.r.add_text(self.spaceque)
            self.spaceque=None
        string=self.fontregister.get_unicode(self.stackmemory.f,c)
        if not self.is_mathmode > 0:
            if self.r:
                self.r.add_text(string)
            else:
                self.add_new_paragraph()
                self.r.add_text(string)
        return(ans)


    def add_to_h(self,b):
        """
        add d to stackmemory.add_to_h.
        """
        super().add_to_h(b)
        if b>0:
            if self.p:
                if b>100000:
                    if not self.is_mathmode > 0:
                        self.add_single_space(" ")

    def add_to_v(self,a):
        """
        add a to stackmemory.add_to_v.
        """
        super().add_to_v(a)
        if a > 0:
            if self.p:
                if not self.is_mathmode > 0:
                    self.add_single_space()

        
    def put(self,c,version,bb):
        """
        function for DVI.put.
        this function calls self.draw_char().
        """
        ans=super().put(c,version,bb)
        self.add_single_space("")

    def xxx(self,k,x,version,bb):
        """
        function for special
        """
        ans=super().xxx(k,x,version,bb)
        if x.startswith("texstructure:"):
            lit=x[13:]
            if lit.startswith("start_par"):
                param=lit[10:]
                if len(self.list_stack)>0:
                    (p,q)=self.list_stack.pop()
                    if p==0 and q>0:
                        self.add_new_paragraph(param)
                        self.list_stack.append((0,q))
                    else:
                        self.list_stack.append((p-1,q))
                else:
                    self.add_new_paragraph(param)

            elif lit.startswith("start_"):
                self.add_new_paragraph()
            elif lit.startswith("begin_math"):
                self.is_mathmode=self.is_mathmode+1
                if self.is_mathmode==1:
                    self.is_display=False
                    self.add_mathimage()
            elif lit=="end_math":
                self.is_mathmode=self.is_mathmode-1
            elif lit=="begin_display":
                self.is_mathmode=self.is_mathmode+1
                if self.is_mathmode==1:
                    self.is_display=True
                    self.add_mathimage()
            elif lit=="end_display":
                self.is_mathmode=self.is_mathmode-1
            elif lit.startswith("begin_list"):
                self.list_stack.append((0,0))
                self.add_new_paragraph()

            elif lit=="end_list":
                self.list_stack.pop()
                self.add_new_paragraph()

            elif lit.startswith("begin_trivlist"):
                self.list_stack.append((0,0))
                self.add_new_paragraph()
    
            elif lit=="end_trivlist":
                self.list_stack.pop()
                self.r.add_break()
                self.add_new_paragraph()

            elif lit=="item":
                (p,q)=self.list_stack.pop()
                if q>0:
                    self.add_new_paragraph()

                q=q+1
                p=1
                self.list_stack.append((p,q))

                    
        return(ans)

    def fnt(self,x,version,bb):
        """
        function for DVI.fnt*
        """
        ans=super().fnt(x,version,bb)
        if not self.is_mathmode > 0:
            self.add_new_run_for_text()
        return(ans)


class PickupMathStackMachine(dviware.DviStackMachine):
    def __init__(self,outfile,debugmode=False):
        super().__init__(debugmode)
        self.outfile = outfile
        self.is_mathmode = 0

    def is_acceptable_font(self,k):
        """
        Now we assume dvipng. So remove japanese.
        """
        if self.fontregister.get_encoding(k)=="JIS":
            return False
        if self.fontregister.get_encoding(k)=="U":
            return False
        return True

    def begin_page(self):
        boppos=self.outfile.tell()
        self.outfile.write(self.bytes_for_bop)
        self.outfile.write(self.previous_bop.to_bytes(4,byteorder='big',signed=True))
        self.previous_bop=boppos
        self.num_of_pages=self.num_of_pages+1
        self.outfile.write(self.queue_of_bytes)
        self.queue_of_bytes=bytes(0)
        if len(self.bytes_for_page)>0:
            self.outfile.write(self.bytes_for_page)
        self.stack_depth=self.stack_depth_for_bytes_for_page

    def end_page(self):
        for i in range(self.stack_depth):
            self.outfile.write(dviware.DVI.pop.to_bytes(1,byteorder='big',signed=False))
        self.outfile.write(dviware.DVI.eop.to_bytes(1,byteorder='big',signed=False))

    def draw_char(self,h,v,c):
        """
        function to draw c.
        """
        pass

    def draw_box(self,h,v,a,b):
        """
        primitive function to draw a box.
        FIXME
        """
        pass

    def set(self,c,version,bb):
        """
        function for DVI.put.
        this function calls self.draw_char().
        FIXME TO MOVE
        """
        ans=super().set(c,version,bb)
        if self.is_mathmode > 0:
            if self.is_acceptable_font(self.stackmemory.f):
                self.outfile.write(bb)

        
    def set_rule(self,a,b,bb):
        """
        function for DVI.set_rule.
        this function calls self.draw_box().
        FIXME TO MOVE
        """
        ans=super().set_rule(a,b,bb)
        if self.is_mathmode > 0:
            self.outfile.write(bb)
    
    def put(self,c,version,bb):
        """
        function for DVI.put.
        this function calls self.draw_char().
        """
        ans=super().put(c,version,bb)
        if self.is_mathmode > 0:
            if self.is_acceptable_font(self.stackmemory.f):
                self.outfile.write(bb)

    def put_rule(self,a,b,bb):
        """
        function for DVI.put_rule
        this function calls self.draw_box().
        """
        ans=super().put_rule(a,b,bb)
        if self.is_mathmode > 0:
            self.outfile.write(bb)

    def nop(self,bb):
        """
        function for DVI.nop.
        """
        ans=super().nop(bb)
        if self.is_mathmode > 0:
            self.outfile.write(bb)
        else:
            self.bytes_for_page=self.bytes_for_page+bb

    def bop(self,cc,p,bb):
        """
        function for DVI.bop.
        FIXME
        """
        ans=super().bop(cc,p,bb)
        self.bytes_for_bop=bb[:-4]
        self.bytes_for_page=bytes(0)
        self.stack_depth=0
        self.stack_depth_for_bytes_for_page=0
        if self.is_mathmode > 0:
            self.begin_page()
                    
    def eop(self,bb):
        """
        function for DVI.eop
        """
        ans=super().eop(bb)
        if self.is_mathmode > 0:
            self.end_page()

    def push(self,bb):
        """
        funtion for stack. 
        function for DVI.push.
        """
        ans=super().push(bb)
        self.bytes_for_page=self.bytes_for_page+bb
        self.stack_depth_for_bytes_for_page=self.stack_depth_for_bytes_for_page+1
        if self.is_mathmode > 0:
            self.outfile.write(bb)
            self.stack_depth=self.stack_depth+1

    def pop(self,bb):
        """
        funtion for stack.
        function for DVI.pop
        """
        self.bytes_for_page=self.bytes_for_page+bb
        self.stack_depth_for_bytes_for_page=self.stack_depth_for_bytes_for_page-1
        if self.is_mathmode > 0:
            self.outfile.write(bb)
            self.stack_depth=self.stack_depth-1

    def right(self,b,version,bb):
        """
        funtion for stack.
        function for DVI.right*
        """
        ans=super().right(b,version,bb)
        self.bytes_for_page=self.bytes_for_page+bb
        if self.is_mathmode > 0:
            self.outfile.write(bb)

    def w(self,b,version,bb):
        """
        funtion for stack.
        function for DVI.w*
        """
        ans=super().w(b,version,bb)
        self.bytes_for_page=self.bytes_for_page+bb
        if self.is_mathmode > 0:
            self.outfile.write(bb)
        
    def x(self,b,version,bb):
        """
        funtion for stack.
        function for DVI.x*
        """
        ans=super().x(b,version,bb)
        self.bytes_for_page=self.bytes_for_page+bb
        if self.is_mathmode > 0:
            self.outfile.write(bb)

    def down(self,a,version,bb):
        """
        funtion for stack.
        function for DVI.down*
        """
        ans=super().down(a,version,bb)
        self.bytes_for_page=self.bytes_for_page+bb
        if self.is_mathmode > 0:
            self.outfile.write(bb)

    def y(self,a,version,bb):
        """
        funtion for stack.
        function for DVI.y*
        """
        ans=super().y(a,version,bb)
        self.bytes_for_page=self.bytes_for_page+bb
        if self.is_mathmode > 0:
            self.outfile.write(bb)
        
    def z(self,a,version,bb):
        """
        funtion for stack.
        function for DVI.z*
        """
        ans=super().z(a,version,bb)
        self.bytes_for_page=self.bytes_for_page+bb
        if self.is_mathmode > 0:
            self.outfile.write(bb)
    
    def fnt(self,x,version,bb):
        """
        function for DVI.fnt*
        """
        ans=super().fnt(x,version,bb)
        if self.is_acceptable_font(x):
            self.bytes_for_page=self.bytes_for_page+bb
            if self.is_mathmode > 0:
                self.outfile.write(bb)
        return ans

        
    def xxx(self,k,x,version,bb):
        """
        function for DVI.xxx*
        function for special
        FIXME
        """
        ans=super().xxx(k,x,version,bb)
        if x.startswith("texstructure:"):
            lit=x[13:]
            if lit.startswith("begin_math"):
                self.is_mathmode=self.is_mathmode+1
                if self.is_mathmode==1:
                    self.begin_page()
            elif lit=="end_math":
                self.is_mathmode=self.is_mathmode-1
                if self.is_mathmode==0:
                    self.end_page()
            elif lit=="begin_display":
                self.is_mathmode=self.is_mathmode+1
                if self.is_mathmode==1:
                    self.begin_page()
            elif lit=="end_display":
                self.is_mathmode=self.is_mathmode-1
                if self.is_mathmode==0:
                    self.end_page()
        else:
            if self.is_mathmode > 0:
                self.outfile.write(bb)
        return(ans)

        
    def fnt_def(self,k,c,s,d,a,l,n,version,bb):
        """
        function for DVI.fnt_def**
        FIXME
        """
        ans=super().fnt_def(k,c,s,d,a,l,n,version,bb)
        if self.is_acceptable_font(k):
            if self.is_mathmode > 0:
                self.outfile.write(bb)
            elif self.is_postemble: 
                self.outfile.write(bb)
            else:
                self.queue_of_bytes=self.queue_of_bytes+bb
        return ans

    def pre(self,i,num,den,mag,k,x,bb):
        """
        function for DVI.pre.
        read preamble and check the version of dvi.
        """
        ans=super().pre(i,num,den,mag,k,x,bb)
        self.outfile.write(bb)
        self.previous_bop=-1
        self.num_of_pages=0
        self.is_postemble=False
        self.is_mathmode=0
        self.queue_of_bytes=bytes(0)

    def post(self,p,num,den,mag,l,u,s,t,bb):
        """
        function for DVI.post.
        read postamble.
        """
        ans=super().post(p,num,den,mag,l,u,s,t,bb)
        self.is_postemble=True
        postpos=self.outfile.tell()
        self.outfile.write(bb[:1])
        self.outfile.write(self.previous_bop.to_bytes(4,byteorder='big',signed=True))
        self.outfile.write(bb[5:-2])
        self.outfile.write(self.num_of_pages.to_bytes(2,byteorder='big',signed=False))
        self.previous_bop=postpos


    def post_post(self,q,i,bb):
        """
        function for DVI.post_post.
        read postamble.
        """
        ans=super().post_post(q,i,bb)
        self.outfile.write(bb[:1])
        self.outfile.write(self.previous_bop.to_bytes(4,byteorder='big',signed=True))
        self.outfile.write(bb[5:6])        
        bbb=(223).to_bytes(1,byteorder='big',signed=False)
        for i in range(4):
            self.outfile.write(bbb)
        while self.outfile.tell() % 4 != 0:
            self.outfile.write(bbb)
        
    def d(self,d,bb):
        """
        function for DVI.d for pDVI format.
        """
        ans=super().d(d,bb)
        self.bytes_for_page=self.bytes_for_page+bb
        if self.is_mathmode > 0:
            self.outfile.write(bb)



def test():
    if len(sys.argv)<2:
        print("usage: python3 "+sys.argv[0]+" dvifile")
        return
    filename=sys.argv[1]
    with open(filename, mode='r+b') as file:
        outfilename=sys.argv[1]+'.math.dvi'
        with open(outfilename, mode='w+b') as outfile:            
            dvistackmachine=PickupMathStackMachine(outfile,debugmode=False)
            dviinterpreter=dviware.DviInterpreter(file,dvistackmachine)
            dviinterpreter.readCodes()
#    for i in range(dvistackmachine.total_pages):
#        os.system("dvipng -T tight -pp "+str(i+1)+" "+outfilename)
    os.system("dvipng -T tight "+outfilename)
    with open(filename, mode='r+b') as file:    
        dvistackmachine=DviDocxStackMachine(sys.argv[1]+'.math',debugmode=False)
        dviinterpreter=dviware.DviInterpreter(file,dvistackmachine)
        dviinterpreter.readCodes()
        dvistackmachine.document.save(sys.argv[1]+'.docx')
        
if __name__=="__main__":
    test()


